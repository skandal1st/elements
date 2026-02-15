"""Сервис работы с .docx шаблонами: парсинг, плейсхолдеры, генерация."""

import os
import re
import uuid
from copy import deepcopy
from pathlib import Path

UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "uploads/tickets")).parent / "documents"


def get_template_html_content(docx_path: str) -> str:
    """Парсит .docx и возвращает HTML-представление с data-атрибутами."""
    from docx import Document

    doc = Document(docx_path)
    html_parts = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        # Экранируем HTML
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Подсвечиваем плейсхолдеры {{...}}
        text = re.sub(
            r"\{\{(\w+)\}\}",
            r'<span class="placeholder" data-key="\1">{{\1}}</span>',
            text,
        )
        style = ""
        if para.style and para.style.name:
            sn = para.style.name.lower()
            if "heading 1" in sn:
                style = ' style="font-size:1.5em;font-weight:bold;"'
            elif "heading 2" in sn:
                style = ' style="font-size:1.25em;font-weight:bold;"'
            elif "heading 3" in sn:
                style = ' style="font-size:1.1em;font-weight:bold;"'
        html_parts.append(f'<p data-paragraph="{idx}"{style}>{text}</p>')
    return "\n".join(html_parts)


def replace_text_with_placeholder(
    docx_path: str,
    paragraph_index: int,
    start: int,
    end: int,
    placeholder_key: str,
) -> None:
    """Заменяет текст в указанном параграфе на маркер {{key}} в .docx файле."""
    from docx import Document

    doc = Document(docx_path)
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise ValueError(f"Параграф с индексом {paragraph_index} не найден")

    para = doc.paragraphs[paragraph_index]
    full_text = para.text
    if start < 0 or end > len(full_text) or start >= end:
        raise ValueError("Неверные границы текста")

    new_text = full_text[:start] + "{{" + placeholder_key + "}}" + full_text[end:]

    # Очищаем runs и записываем новый текст
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    doc.save(docx_path)


def generate_document_from_template(
    docx_path: str, values: dict
) -> dict:
    """Заменяет все {{key}} на значения и сохраняет как новый .docx файл."""
    from docx import Document

    doc = Document(docx_path)

    for para in doc.paragraphs:
        for run in para.runs:
            if "{{" in run.text:
                text = run.text
                for key, value in values.items():
                    text = text.replace("{{" + key + "}}", str(value))
                run.text = text

    # Также обработаем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "{{" in run.text:
                            text = run.text
                            for key, value in values.items():
                                text = text.replace("{{" + key + "}}", str(value))
                            run.text = text

    # Сохраняем новый файл
    dest_dir = UPLOAD_DIR / "files"
    dest_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.docx"
    dest = dest_dir / filename
    doc.save(str(dest))

    return {
        "file_path": f"/uploads/documents/files/{filename}",
        "file_name": filename,
        "file_size": dest.stat().st_size,
    }
